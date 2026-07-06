"""
Internal helper functions for table manipulation in Word templates.

This module contains low-level functions for filling Word tables with DataFrame data,
including support for MultiIndex DataFrames, merged cells, and flexible styling.

Private module - Not intended for direct external use.
Import from the public API in api.py instead.
"""

import pandas as pd
import re
from typing import Union, Optional, List, Tuple
from docx.oxml.ns import qn
from .schemas_helpers import EstilosTabla, OpcionesTabla
from ._table_styling import apply_cell_style, apply_row_height, apply_table_borders


def insertar_titulo_de_tabla_con_bookmark(doc, marcador_tabla, titulo, bookmark, estilo_titulo="Normal"):
    """Inserta un título con bookmark antes de la tabla que contiene el marcador."""
    if not titulo:
        return

    tabla_objetivo = None
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                if marcador_tabla in celda.text:
                    tabla_objetivo = tabla
                    break
            if tabla_objetivo is not None:
                break
        if tabla_objetivo is not None:
            break

    if tabla_objetivo is None:
        raise ValueError(f"No se encontró la tabla con marcador '{marcador_tabla}' para insertar título.")

    from docx.oxml import OxmlElement

    bookmark_id = str(abs(hash(bookmark)) % 1000000)

    nuevo_p = OxmlElement('w:p')

    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), estilo_titulo)
    pPr.append(pStyle)
    nuevo_p.append(pPr)

    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), bookmark_id)
    bookmark_start.set(qn('w:name'), bookmark)
    nuevo_p.append(bookmark_start)

    run = OxmlElement('w:r')
    text = OxmlElement('w:t')
    text.set(qn('xml:space'), 'preserve')
    text.text = str(titulo)
    run.append(text)
    nuevo_p.append(run)

    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), bookmark_id)
    nuevo_p.append(bookmark_end)

    tabla_objetivo._element.addprevious(nuevo_p)


# ===== NORMALIZATION FUNCTIONS =====

def _normalize_input_to_dataframe(datos: Union[pd.DataFrame, dict]) -> pd.DataFrame:
    """
    Normaliza la entrada a DataFrame.
    
    Args:
        datos: DataFrame o diccionario de columnas
    
    Returns:
        DataFrame normalizado
    
    Raises:
        ValueError: Si los datos están vacíos o no son válidos
    """
    if isinstance(datos, pd.DataFrame):
        df = datos
    elif isinstance(datos, dict):
        df = pd.DataFrame(datos)
    else:
        raise ValueError(
            f"datos debe ser DataFrame o diccionario. Recibido: {type(datos)}"
        )
    
    if df.empty:
        raise ValueError("El DataFrame está vacío. No hay datos para insertar.")
    
    return df


def _normalize_config_estilos(config_estilos: Union[EstilosTabla, dict, None], doc) -> dict:
    """
    Normaliza la configuración de estilos a diccionario.
    
    Args:
        config_estilos: Objeto EstilosTabla, dict, o None
        doc: Objeto Document (para crear EstilosTabla por defecto si es None)
    
    Returns:
        Diccionario de configuración normalizado
    """
    if config_estilos is None:
        # Crear EstilosTabla por defecto
        config_obj = EstilosTabla()
        return config_obj.to_dict()
    elif isinstance(config_estilos, EstilosTabla):
        return config_estilos.to_dict()
    elif isinstance(config_estilos, dict):
        # Validar estructura básica del diccionario
        required_keys = ["por_defecto", "por_columna", "por_fila", "por_celda"]
        if not all(key in config_estilos for key in required_keys):
            raise ValueError(
                f"Diccionario config_estilos debe contener las claves: {required_keys}"
            )
        # Compatibilidad: agregar por_header si no está presente (diccionarios viejos)
        if "por_header" not in config_estilos:
            config_estilos = {**config_estilos, "por_header": {}}
        return config_estilos
    else:
        raise ValueError(
            f"config_estilos debe ser EstilosTabla, dict o None. Recibido: {type(config_estilos)}"
        )


def _normalize_opciones_tabla(opciones_tabla: Union[OpcionesTabla, dict, None]) -> dict:
    """
    Normaliza las opciones de tabla a diccionario.
    
    Args:
        opciones_tabla: Objeto OpcionesTabla, dict, o None
    
    Returns:
        Diccionario de opciones normalizado
    """
    if opciones_tabla is None:
        # Crear OpcionesTabla por defecto
        opciones_obj = OpcionesTabla()
        return opciones_obj.to_dict()
    elif isinstance(opciones_tabla, OpcionesTabla):
        return opciones_tabla.to_dict()
    elif isinstance(opciones_tabla, dict):
        # Agregar valores por defecto para claves faltantes
        defaults = OpcionesTabla().to_dict()
        return {**defaults, **opciones_tabla}
    else:
        raise ValueError(
            f"opciones_tabla debe ser OpcionesTabla, dict o None. Recibido: {type(opciones_tabla)}"
        )


# ===== TABLE SEARCH FUNCTIONS =====

def _find_table_with_marker(doc, nombre_marcador: str) -> Tuple[object, int, int]:
    """
    Busca una tabla que contenga el marcador especificado.
    
    Args:
        doc: Objeto Document
        nombre_marcador: Marcador a buscar (ej: "<<table_sondas>>")
    
    Returns:
        Tupla (tabla, fila_marcador, columna_marcador)
    
    Raises:
        ValueError: Si no se encuentra el marcador
    """
    for table in doc.tables:
        # Buscar en todas las celdas de la tabla
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                if nombre_marcador in cell.text:
                    return (table, row_idx, col_idx)
    
    raise ValueError(
        f"No se encontró el marcador '{nombre_marcador}' en ninguna tabla del documento."
    )


def _disable_row_header_repeat(row):
    """
    Desactiva la propiedad de repetición de encabezado en una fila.
    
    En Word, las filas de encabezado tienen la propiedad tblHeader que hace
    que se repitan en cada salto de página. Esta función elimina esa propiedad.
    
    Args:
        row: Objeto Row de python-docx
    """
    tr = row._element
    trPr = tr.get_or_add_trPr()
    
    # Buscar y eliminar el elemento tblHeader si existe
    tblHeader = trPr.find(qn('w:tblHeader'))
    if tblHeader is not None:
        trPr.remove(tblHeader)


def _delete_row(table, row_idx: int):
    """
    Elimina una fila de la tabla manipulando el XML directamente.
    
    python-docx no proporciona una API directa para eliminar filas,
    por lo que se debe manipular el XML.
    
    Args:
        table: Objeto Table de python-docx
        row_idx: Índice de la fila a eliminar (0-based)
    """
    tbl = table._element
    tr = table.rows[row_idx]._element
    tbl.remove(tr)


# ===== STYLE MERGING FUNCTIONS =====

def _merge_style_config(config: dict, row_idx: int, col_idx: int) -> dict:
    """
    Combina la configuración de estilos para una celda específica.
    
    Resuelve la jerarquía: celda > fila > columna > defecto
    
    Args:
        config: Diccionario de configuración completo
        row_idx: Índice de fila (0-based)
        col_idx: Índice de columna (0-based)
    
    Returns:
        Diccionario con configuración mergeada para esta celda
    """
    merged = {}
    
    # 1. Aplicar configuración por defecto
    if "por_defecto" in config:
        merged.update(config["por_defecto"])
    
    # 2. Aplicar configuración de columna (sobrescribe defecto)
    if "por_columna" in config and col_idx in config["por_columna"]:
        merged.update(config["por_columna"][col_idx])
    
    # 3. Aplicar configuración de fila (sobrescribe columna y defecto)
    if "por_fila" in config and row_idx in config["por_fila"]:
        merged.update(config["por_fila"][row_idx])
    
    # 4. Aplicar configuración de celda (sobrescribe todo)
    if "por_celda" in config and (row_idx, col_idx) in config["por_celda"]:
        merged.update(config["por_celda"][(row_idx, col_idx)])
    
    return merged


# ===== MULTIINDEX FLATTENING FUNCTIONS =====

def _flatten_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    """
    Aplana un DataFrame con MultiIndex en filas o columnas.
    
    Args:
        df: DataFrame potencialmente con MultiIndex
    
    Returns:
        DataFrame aplanado con índice simple
    """
    df_flat = df.copy()
    
    # Aplanar MultiIndex en filas (índice)
    if isinstance(df_flat.index, pd.MultiIndex):
        df_flat = df_flat.reset_index()
    
    # Aplanar MultiIndex en columnas
    if isinstance(df_flat.columns, pd.MultiIndex):
        # Unir niveles con '_'
        df_flat.columns = ['_'.join(map(str, col)).strip('_') for col in df_flat.columns]
    
    return df_flat


# ===== MERGED CELLS FUNCTIONS =====

def _detect_merge_regions_vertical(
    df: pd.DataFrame,
    columnas_para_merge: Optional[List[int]]
) -> List[Tuple[int, int, int]]:
    """
    Detecta regiones de celdas a combinar verticalmente.
    
    Busca valores consecutivos repetidos en columnas especificadas.
    
    Args:
        df: DataFrame con los datos
        columnas_para_merge: Lista de índices de columnas o None (todas)
    
    Returns:
        Lista de tuplas (fila_inicio, fila_fin, columna)
    """
    regions = []
    
    # Determinar qué columnas analizar
    if columnas_para_merge is None:
        cols_to_check = range(len(df.columns))
    else:
        cols_to_check = columnas_para_merge
    
    for col_idx in cols_to_check:
        if col_idx >= len(df.columns):
            continue  # Saltar si el índice está fuera de rango
        
        col_name = df.columns[col_idx]
        values = df[col_name].tolist()
        
        # Detectar secuencias de valores repetidos
        start_idx = 0
        while start_idx < len(values):
            current_value = values[start_idx]
            end_idx = start_idx
            
            # Contar cuántas filas consecutivas tienen el mismo valor
            while end_idx + 1 < len(values) and values[end_idx + 1] == current_value:
                end_idx += 1
            
            # Si hay al menos 2 celdas consecutivas con el mismo valor, agregar región
            if end_idx > start_idx:
                regions.append((start_idx, end_idx, col_idx))
            
            start_idx = end_idx + 1
    
    return regions


def _apply_merged_cells_vertical(table, regions: List[Tuple[int, int, int]], fila_inicio_datos: int):
    """
    Aplica merged cells verticales a la tabla.
    
    Args:
        table: Objeto Table de python-docx
        regions: Lista de tuplas (fila_inicio, fila_fin, columna)
        fila_inicio_datos: Índice de la fila donde inician los datos (después del marcador)
    """
    for start_row, end_row, col in regions:
        # Ajustar índices de fila para la tabla real (sumar fila_inicio_datos)
        table_start_row = fila_inicio_datos + start_row
        table_end_row = fila_inicio_datos + end_row
        
        try:
            # Obtener celdas a combinar
            cell_start = table.cell(table_start_row, col)
            cell_end = table.cell(table_end_row, col)
            
            # Combinar celdas
            cell_start.merge(cell_end)
        except Exception as e:
            # Si falla el merge, continuar con las demás (no es crítico)
            print(f"Advertencia: No se pudo combinar celdas ({start_row}-{end_row}, {col}): {e}")


# ===== MAIN FILL TABLE FUNCTION =====

def fill_table(
    doc,
    nombre_marcador: str,
    datos: Union[pd.DataFrame, dict],
    config_estilos: Union[EstilosTabla, dict, None] = None,
    opciones_tabla: Union[OpcionesTabla, dict, None] = None
):
    """
    Rellena una tabla dinámicamente buscando un marcador e insertando datos de un DataFrame.
    
    Esta es la función principal refactorizada que orquesta todo el proceso.
    
    Args:
        doc: Objeto Document de python-docx
        nombre_marcador: Marcador a buscar en la tabla (ej: "<<table_sondas>>")
        datos: DataFrame o diccionario con los datos a insertar
        config_estilos: Configuración de estilos (objeto EstilosTabla, dict, o None)
        opciones_tabla: Opciones de tabla (objeto OpcionesTabla, dict, o None)
    
    Returns:
        Objeto Document modificado
    
    Raises:
        ValueError: Si hay errores en la configuración o datos
    
    Flujo:
        1. Normalizar inputs (DataFrame, config, opciones)
        2. Buscar tabla con marcador
        3. Aplanar DataFrame si tiene MultiIndex
        4. Insertar datos fila por fila
        5. Detectar y aplicar merged cells (si está activado)
        6. Aplicar estilos a cada celda
        7. Eliminar fila marcador (si está activado)
    """
    # ===== PASO 1: NORMALIZAR INPUTS =====
    df = _normalize_input_to_dataframe(datos)
    config = _normalize_config_estilos(config_estilos, doc)
    opciones = _normalize_opciones_tabla(opciones_tabla)
    
    # ===== PASO 2: BUSCAR TABLA CON MARCADOR =====
    table, fila_marcador, col_marcador = _find_table_with_marker(doc, nombre_marcador)
    
    # ===== BUG FIX 1: DESACTIVAR REPETICIÓN DE ENCABEZADO EN FILA MARCADOR =====
    # La fila del marcador puede tener la propiedad tblHeader (encabezado repetible)
    # que hace que se repita en cada salto de página. Debemos desactivarla.
    _disable_row_header_repeat(table.rows[fila_marcador])
    
    # ===== PASO 3: APLANAR DATAFRAME SI TIENE MULTIINDEX =====
    if opciones["aplanar_multiindex"]:
        if isinstance(df.index, pd.MultiIndex) or isinstance(df.columns, pd.MultiIndex):
            df = _flatten_dataframe(df)
    else:
        # Si no se debe aplanar y tiene MultiIndex, lanzar error
        if isinstance(df.index, pd.MultiIndex) or isinstance(df.columns, pd.MultiIndex):
            raise ValueError(
                "El DataFrame tiene MultiIndex pero 'aplanar_multiindex' está desactivado. "
                "Active la opción o aplane el DataFrame manualmente con df.reset_index()."
            )
    
    # ===== PASO 4: VALIDAR QUE HAY SUFICIENTES COLUMNAS =====
    num_cols_df = len(df.columns)
    num_cols_table = len(table.columns)
    if num_cols_df > num_cols_table:
        raise ValueError(
            f"El DataFrame tiene {num_cols_df} columnas pero la tabla solo tiene {num_cols_table}. "
            "Ajuste la plantilla o el DataFrame."
        )
    
    # ===== PASO 5: AGREGAR FILAS NECESARIAS =====
    # Calcular cuántas filas necesitamos agregar
    # Si eliminamos el marcador: usamos esa fila + las que siguen
    # Si NO eliminamos: necesitamos todas las filas de datos desde el marcador
    filas_necesarias = len(df)
    filas_actuales_disponibles = len(table.rows) - fila_marcador  # Incluye fila marcador
    
    if filas_necesarias > filas_actuales_disponibles:
        # Agregar solo las filas faltantes (BUG FIX 2: cálculo corregido)
        filas_a_agregar = filas_necesarias - filas_actuales_disponibles
        for _ in range(filas_a_agregar):
            table.add_row()
    
    # ===== PASO 6: DETECTAR REGIONES DE MERGE (antes de insertar datos) =====
    # BUG FIX 3: Detectar merge ANTES de insertar texto para evitar duplicados
    merge_regions = []
    if opciones["detectar_merge"]:
        merge_regions = _detect_merge_regions_vertical(df, opciones["columnas_para_merge"])
    
    # Crear un set de celdas que NO deben recibir texto (segunda celda en adelante de merge)
    celdas_merged_secundarias = set()
    for start_row, end_row, col in merge_regions:
        # Agregar todas las celdas excepto la primera de cada región
        for row_idx in range(start_row + 1, end_row + 1):
            celdas_merged_secundarias.add((row_idx, col))
    
    # ===== PASO 7: INSERTAR DATOS FILA POR FILA =====
    # Siempre empezamos a escribir desde la fila del marcador
    fila_inicio_datos = fila_marcador
    
    for df_row_idx, (_, df_row) in enumerate(df.iterrows()):
        # Índice de fila en la tabla (después del marcador)
        table_row_idx = fila_inicio_datos + df_row_idx
        
        # Insertar datos en cada columna
        for df_col_idx, col_name in enumerate(df.columns):
            # BUG FIX 3: No escribir en celdas secundarias de regiones mergeadas
            if (df_row_idx, df_col_idx) not in celdas_merged_secundarias:
                cell = table.cell(table_row_idx, df_col_idx)
                cell.text = str(df_row[col_name])
    
    # ===== PASO 8: APLICAR MERGED CELLS =====
    if opciones["detectar_merge"]:
        _apply_merged_cells_vertical(table, merge_regions, fila_inicio_datos)
    
    # ===== PASO 9a: APLICAR ESTILOS A FILAS DE HEADER =====
    # por_header usa índices explícitos de filas de la tabla (0-based, desde el tope).
    # Esto permite estilizar headers aunque el marcador esté en la fila 0.
    # También soporta claves string (por ejemplo "0") por compatibilidad con dicts serializados.
    config_header = config.get("por_header", {})
    if config_header:
        for raw_idx, header_config in config_header.items():
            if not header_config:
                continue

            try:
                header_row_idx = int(raw_idx)
            except (TypeError, ValueError):
                continue

            if header_row_idx < 0 or header_row_idx >= len(table.rows):
                continue

            # Combinar por_defecto con header_config (igual que _merge_style_config)
            merged_header_config = {}
            if "por_defecto" in config:
                merged_header_config.update(config["por_defecto"])
            merged_header_config.update(header_config)

            for col_idx in range(len(table.columns)):
                cell = table.cell(header_row_idx, col_idx)
                apply_cell_style(cell, doc, config, header_row_idx, col_idx, merged_header_config)

    # ===== PASO 9: APLICAR ESTILOS A CADA CELDA =====
    for df_row_idx in range(len(df)):
        table_row_idx = fila_inicio_datos + df_row_idx
        row = table.rows[table_row_idx]
        
        # Aplicar altura de fila
        altura_fila = config["por_defecto"].get("altura_fila")
        if altura_fila:
            apply_row_height(row, altura_fila)
        
        # Aplicar estilos a cada celda
        # BUG FIX 3: Solo aplicar estilos a celdas que no son secundarias en merge
        for df_col_idx in range(len(df.columns)):
            if (df_row_idx, df_col_idx) not in celdas_merged_secundarias:
                cell = table.cell(table_row_idx, df_col_idx)
                merged_config = _merge_style_config(config, df_row_idx, df_col_idx)
                apply_cell_style(cell, doc, config, df_row_idx, df_col_idx, merged_config)
    
    # ===== PASO 10: ELIMINAR FILAS SOBRANTES AL FINAL =====
    # BUG FIX 2: Eliminar filas vacías que puedan quedar al final de la tabla
    # Esto puede ocurrir si la plantilla tiene filas preexistentes o si add_row()
    # creó filas de más
    ultima_fila_con_datos = fila_inicio_datos + len(df) - 1
    total_filas = len(table.rows)
    
    # Eliminar todas las filas después de la última fila con datos
    # Nota: Eliminamos de atrás hacia adelante para evitar problemas de índices
    for row_idx in range(total_filas - 1, ultima_fila_con_datos, -1):
        _delete_row(table, row_idx)
    
    # Aplicar bordes a la tabla
    apply_table_borders(table)
    
    return doc


def _remove_paragraph(paragraph):
    """Elimina un parrafo del documento manipulando su XML."""
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


def _find_paragraph_with_marker(doc, marcador: str):
    """Busca el primer parrafo que contiene un marcador."""
    for paragraph in doc.paragraphs:
        if marcador in paragraph.text:
            return paragraph
    raise ValueError(f"No se encontro el marcador '{marcador}' en los parrafos del documento.")


def _replace_marker_in_paragraph(paragraph, marcador: str, replacement: str = ""):
    """Reemplaza un marcador en el texto consolidado del parrafo."""
    full_text = "".join(run.text for run in paragraph.runs)
    if marcador not in full_text:
        return

    updated_text = full_text.replace(marcador, replacement)
    first_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    first_run.text = updated_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _insert_table_after_paragraph(doc, paragraph, rows: int, cols: int):
    """Crea una tabla y la reubica despues del parrafo indicado."""
    table = doc.add_table(rows=rows, cols=cols)
    paragraph._p.addnext(table._tbl)
    return table


def _insert_title_before_table(table, titulo: str, bookmark: str, estilo_titulo: str = "Normal"):
    """Inserta un caption de tabla numerado con bookmark antes de una tabla."""
    if not titulo or not bookmark:
        return

    from docx.oxml import OxmlElement

    bookmark_id = str(abs(hash(bookmark)) % 1000000)
    titulo_normalizado = _normalizar_texto_de_titulo_tabla(titulo)

    nuevo_p = OxmlElement('w:p')

    p_pr = OxmlElement('w:pPr')
    p_style = OxmlElement('w:pStyle')
    p_style.set(qn('w:val'), estilo_titulo)
    p_pr.append(p_style)
    nuevo_p.append(p_pr)

    run_label = OxmlElement('w:r')
    text_label = OxmlElement('w:t')
    text_label.set(qn('xml:space'), 'preserve')
    text_label.text = 'Tabla '
    run_label.append(text_label)
    nuevo_p.append(run_label)

    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), bookmark_id)
    bookmark_start.set(qn('w:name'), bookmark)
    nuevo_p.append(bookmark_start)

    run_begin = OxmlElement('w:r')
    fldchar_begin = OxmlElement('w:fldChar')
    fldchar_begin.set(qn('w:fldCharType'), 'begin')
    run_begin.append(fldchar_begin)
    nuevo_p.append(run_begin)

    run_instr = OxmlElement('w:r')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' SEQ Tabla \\* ARABIC '
    run_instr.append(instr_text)
    nuevo_p.append(run_instr)

    run_separate = OxmlElement('w:r')
    fldchar_separate = OxmlElement('w:fldChar')
    fldchar_separate.set(qn('w:fldCharType'), 'separate')
    run_separate.append(fldchar_separate)
    nuevo_p.append(run_separate)

    run_result = OxmlElement('w:r')
    text_result = OxmlElement('w:t')
    text_result.text = '1'
    run_result.append(text_result)
    nuevo_p.append(run_result)

    run_end = OxmlElement('w:r')
    fldchar_end = OxmlElement('w:fldChar')
    fldchar_end.set(qn('w:fldCharType'), 'end')
    run_end.append(fldchar_end)
    nuevo_p.append(run_end)

    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), bookmark_id)
    nuevo_p.append(bookmark_end)

    run_suffix = OxmlElement('w:r')
    text_suffix = OxmlElement('w:t')
    text_suffix.set(qn('xml:space'), 'preserve')
    text_suffix.text = f'. {titulo_normalizado}'
    run_suffix.append(text_suffix)
    nuevo_p.append(run_suffix)

    table._element.addprevious(nuevo_p)


def _normalizar_texto_de_titulo_tabla(titulo: str) -> str:
    """Elimina prefijos manuales tipo 'Tabla 1.' para evitar duplicar el caption."""
    titulo_limpio = str(titulo).strip()
    patron = re.compile(r'^Tabla\s+[A-Za-z0-9IVXivx]+\s*[\.:\-]?\s*')
    return patron.sub('', titulo_limpio)


def _crear_bookmark_de_tabla_desde_variable(variable: str) -> str:
    """Genera un bookmark RefTabla_* a partir de <<nuevatabla_*>>."""
    nombre_base = variable.strip('<>')
    nombre_base = nombre_base.replace('nuevatabla_', '', 1)
    return f'RefTabla_{nombre_base}'


def crear_tabla_desde_marcador(
    doc,
    marcador_tabla: str,
    datos: Union[pd.DataFrame, dict],
    config_estilos: Union[EstilosTabla, dict, None] = None,
    opciones_tabla: Union[OpcionesTabla, dict, None] = None,
    titulo: str = "",
    bookmark: str = "",
    estilo_titulo: str = "Normal",
):
    """
    Crea una tabla nueva desde cero a partir de un marcador en un parrafo.

    La tabla resultante incluye encabezados (columnas del DataFrame) y contenido.
    """
    df = _normalize_input_to_dataframe(datos)
    config = _normalize_config_estilos(config_estilos, doc)
    opciones = _normalize_opciones_tabla(opciones_tabla)

    if opciones["aplanar_multiindex"]:
        if isinstance(df.index, pd.MultiIndex) or isinstance(df.columns, pd.MultiIndex):
            df = _flatten_dataframe(df)
    else:
        if isinstance(df.index, pd.MultiIndex) or isinstance(df.columns, pd.MultiIndex):
            raise ValueError(
                "El DataFrame tiene MultiIndex pero 'aplanar_multiindex' esta desactivado. "
                "Active la opcion o aplane el DataFrame manualmente con df.reset_index()."
            )

    paragraph = _find_paragraph_with_marker(doc, marcador_tabla)
    table = _insert_table_after_paragraph(doc, paragraph, rows=len(df) + 1, cols=len(df.columns))

    if titulo:
        _insert_title_before_table(table, titulo=titulo, bookmark=bookmark, estilo_titulo=estilo_titulo)

    # Encabezados
    for col_idx, col_name in enumerate(df.columns):
        header_cell = table.cell(0, col_idx)
        header_cell.text = str(col_name)
        if header_cell.paragraphs and header_cell.paragraphs[0].runs:
            header_cell.paragraphs[0].runs[0].bold = True

        # Combinar por_defecto con por_header[0] si existe
        merged_header_config = {}
        if "por_defecto" in config:
            merged_header_config.update(config["por_defecto"])
        
        # Aplicar configuración específica de header si existe (soporta int 0 o string "0")
        if "por_header" in config:
            for header_key in [0, "0"]:
                if header_key in config["por_header"]:
                    merged_header_config.update(config["por_header"][header_key])
                    break
        
        apply_cell_style(header_cell, doc, config, 0, col_idx, merged_header_config)

    # Detectar merges para el cuerpo de la tabla
    merge_regions = []
    if opciones["detectar_merge"]:
        merge_regions = _detect_merge_regions_vertical(df, opciones["columnas_para_merge"])

    celdas_merged_secundarias = set()
    for start_row, end_row, col in merge_regions:
        for row_idx in range(start_row + 1, end_row + 1):
            celdas_merged_secundarias.add((row_idx, col))

    # Datos (comienzan en la fila 1 de la tabla)
    for df_row_idx, (_, df_row) in enumerate(df.iterrows()):
        table_row_idx = df_row_idx + 1

        row = table.rows[table_row_idx]
        altura_fila = config["por_defecto"].get("altura_fila")
        if altura_fila:
            apply_row_height(row, altura_fila)

        for df_col_idx, col_name in enumerate(df.columns):
            if (df_row_idx, df_col_idx) not in celdas_merged_secundarias:
                cell = table.cell(table_row_idx, df_col_idx)
                cell.text = str(df_row[col_name])
                merged_config = _merge_style_config(config, df_row_idx, df_col_idx)
                apply_cell_style(cell, doc, config, df_row_idx, df_col_idx, merged_config)

    if opciones["detectar_merge"]:
        _apply_merged_cells_vertical(table, merge_regions, fila_inicio_datos=1)

    # Aplicar bordes a la tabla
    apply_table_borders(table)

    _replace_marker_in_paragraph(paragraph, marcador_tabla, "")
    if not paragraph.text.strip():
        _remove_paragraph(paragraph)

    return doc


def procesar_reemplazar_variable_por_tabla(doc, diccionario_de_reemplazos: dict):
    """Orquesta la creacion de tablas <<nuevatabla_*>> y prepara <<refnuevatabla_*>>."""
    if diccionario_de_reemplazos is None:
        raise ValueError("El diccionario de reemplazos no puede ser None.")

    variables_tabla = {
        k: v for k, v in diccionario_de_reemplazos.items()
        if k.startswith("<<nuevatabla_")
    }

    referencias_tablas = {}

    for variable, config_tabla in variables_tabla.items():
        if not isinstance(config_tabla, dict):
            raise ValueError(
                f"La variable '{variable}' debe contener un diccionario con configuracion de tabla."
            )

        keys_requeridas = ["tabla", "titulo"]
        faltantes = [key for key in keys_requeridas if key not in config_tabla]
        if faltantes:
            raise ValueError(
                f"La variable '{variable}' no tiene las keys requeridas: {faltantes}."
            )

        tabla_df = config_tabla["tabla"]
        estilos_de_tabla = config_tabla.get("estilos_de_tabla")
        opciones_tabla = config_tabla.get("opciones_de_tabla")
        titulo_tabla = config_tabla["titulo"]
        bookmark_tabla = config_tabla.get("bookmark") or _crear_bookmark_de_tabla_desde_variable(variable)

        if not isinstance(bookmark_tabla, str) or not bookmark_tabla.startswith("RefTabla"):
            raise ValueError(
                f"El bookmark de '{variable}' debe comenzar con 'RefTabla'."
            )

        crear_tabla_desde_marcador(
            doc,
            marcador_tabla=variable,
            datos=tabla_df,
            config_estilos=estilos_de_tabla,
            opciones_tabla=opciones_tabla,
            titulo=titulo_tabla,
            bookmark=bookmark_tabla,
            estilo_titulo="Caption",
        )

        variable_ref_key = variable.replace("<<nuevatabla_", "<<refnuevatabla_", 1)
        referencias_tablas[variable_ref_key] = [bookmark_tabla]

    diccionario_de_reemplazos.update(referencias_tablas)
    return doc


def procesar_rellenar_tablas_en_plantilla(doc, diccionario_de_reemplazos: dict):
    """Orquesta el rellenado de tablas existentes marcadas con <<editartabla...>>."""
    if diccionario_de_reemplazos is None:
        raise ValueError("El diccionario de reemplazos no puede ser None.")

    variables_tabla = {
        k: v for k, v in diccionario_de_reemplazos.items()
        if k.startswith("<<editartabla")
    }

    for variable, config_tabla in variables_tabla.items():
        if not isinstance(config_tabla, dict):
            raise ValueError(
                f"La variable '{variable}' debe contener un diccionario con configuracion de tabla."
            )

        if "tabla" not in config_tabla:
            raise ValueError(
                f"La variable '{variable}' debe incluir la key 'tabla'."
            )

        tabla_df = config_tabla["tabla"]
        config_estilos = config_tabla.get("estilos_de_tabla")
        opciones_tabla = config_tabla.get("opciones_de_tabla")

        fill_table(doc, variable, tabla_df, config_estilos, opciones_tabla)

    return doc


def reemplazar_variables_en_tablas_del_documento(doc, diccionario_de_reemplazos):
    """Reemplaza variables de texto dentro de celdas de tablas del documento."""
    from ._text_helpers import replace_text_variables_in_paragraph

    prefijos_excluidos = [
        "<<fig",
        "<<reffigura",
        "<<reftabla_",
        "<<refnuevatabla_",
        "<<nuevatabla_",
        "<<editartabla",
        "<<external_doc",
    ]
    variables_texto = {
        k: v for k, v in diccionario_de_reemplazos.items()
        if not any(k.startswith(prefix) for prefix in prefijos_excluidos)
    }

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    variables_en_parrafo = []
                    for variable, dato in variables_texto.items():
                        if variable in paragraph.text:
                            variables_en_parrafo.append((variable, dato))

                    if variables_en_parrafo:
                        replace_text_variables_in_paragraph(paragraph, variables_en_parrafo)

    return doc

