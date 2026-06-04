"""
Internal helper functions for applying cell styles in Word tables.

This module contains low-level functions for styling table cells including
background colors, text colors, alignment, and paragraph styles.

Private module - Not intended for direct external use.
Import from the public API in api.py instead.
"""

from typing import Optional, Tuple
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.parser import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor


def rgb_to_hex(rgb: Tuple[int, int, int]) -> str:
    """
    Convierte una tupla RGB a formato hexadecimal para XML.
    
    Args:
        rgb: Tupla (R, G, B) con valores 0-255
    
    Returns:
        String hexadecimal sin el prefijo '#' (ej: "E6E6FA")
    
    Example:
        >>> rgb_to_hex((230, 230, 250))
        'E6E6FA'
    """
    return '{:02X}{:02X}{:02X}'.format(rgb[0], rgb[1], rgb[2])


def apply_cell_background_color(cell, color: Optional[Tuple[int, int, int]]):
    """
    Aplica color de fondo a una celda usando XML (python-docx no tiene API directa).
    
    Args:
        cell: Objeto Cell de python-docx
        color: Tupla (R, G, B) con valores 0-255, o None para sin color
    
    Note:
        Esta función manipula directamente el XML del documento porque python-docx
        no proporciona una API de alto nivel para color de fondo de celdas.
    
    Example:
        >>> table = doc.tables[0]
        >>> cell = table.cell(0, 0)
        >>> apply_cell_background_color(cell, (230, 230, 250))  # Lavanda
    """
    if color is None:
        return
    
    # Convertir RGB a hex
    color_hex = rgb_to_hex(color)
    
    # Obtener o crear el elemento tcPr (table cell properties)
    tc_pr = cell._element.get_or_add_tcPr()
    
    # Crear elemento de shading (sombreado/color de fondo)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    
    # Remover shading anterior si existe
    existing_shading = tc_pr.find(qn('w:shd'))
    if existing_shading is not None:
        tc_pr.remove(existing_shading)
    
    # Agregar el nuevo shading
    tc_pr.append(shading_elm)


def apply_cell_vertical_alignment(cell, alignment: str):
    """
    Aplica alineación vertical a una celda.
    
    Args:
        cell: Objeto Cell de python-docx
        alignment: 'center', 'top', o 'bottom'
    
    Note:
        Se reaplica después de aplicar estilos porque hay un bug conocido
        donde aplicar estilos puede "dañar" la alineación vertical.
    """
    alignment_map = {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'top': WD_ALIGN_PARAGRAPH.LEFT,  # En vertical: LEFT = TOP
        'bottom': WD_ALIGN_PARAGRAPH.RIGHT  # En vertical: RIGHT = BOTTOM
    }
    
    if alignment in alignment_map:
        cell.vertical_alignment = alignment_map[alignment]


def apply_cell_paragraph_style(cell, doc, style_name: str):
    """
    Aplica un estilo de párrafo a la celda.
    
    Args:
        cell: Objeto Cell de python-docx
        doc: Objeto Document (para acceder a estilos)
        style_name: Nombre del estilo a aplicar
    
    Note:
        Solo se aplica al primer párrafo de la celda.
    """
    if not style_name:
        return
    
    try:
        # Aplicar estilo al primer párrafo de la celda
        if len(cell.paragraphs) > 0:
            cell.paragraphs[0].style = doc.styles[style_name]
    except KeyError:
        # Si el estilo no existe, ignorar silenciosamente
        # (ya fue validado en EstilosTabla)
        pass


def apply_cell_horizontal_alignment(cell, alignment: str):
    """
    Aplica alineación horizontal al texto de la celda.
    
    Args:
        cell: Objeto Cell de python-docx
        alignment: 'left', 'center', 'right', o 'justify'
    """
    alignment_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    
    if alignment in alignment_map and len(cell.paragraphs) > 0:
        cell.paragraphs[0].alignment = alignment_map[alignment]


def apply_cell_text_color(cell, color: Tuple[int, int, int]):
    """
    Aplica color de texto a una celda.
    
    Args:
        cell: Objeto Cell de python-docx
        color: Tupla (R, G, B) con valores 0-255
    
    Note:
        Aplica el color a todos los runs del primer párrafo.
    """
    if len(cell.paragraphs) > 0:
        rgb_color = RGBColor(color[0], color[1], color[2])
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = rgb_color


def apply_row_height(row, height: int):
    """
    Establece la altura de una fila.
    
    Args:
        row: Objeto Row de python-docx
        height: Altura en twips (288290 ≈ 4cm)
    """
    if height:
        row.height = height


def apply_cell_style(cell, doc, config: dict, row_idx: int, col_idx: int, merged_config: dict):
    """
    Aplica todos los estilos configurados a una celda.
    
    Esta es la función principal que orquesta la aplicación de todos los estilos.
    Resuelve la jerarquía: celda > fila > columna > defecto.
    
    Args:
        cell: Objeto Cell de python-docx
        doc: Objeto Document
        config: Diccionario de configuración de estilos completo
        row_idx: Índice de la fila (0-based)
        col_idx: Índice de la columna (0-based)
        merged_config: Configuración ya mergeada para esta celda específica
    
    Note:
        El parámetro merged_config es el resultado de _merge_style_config()
        que ya ha resuelto la jerarquía. Se mantiene config por si se necesita
        acceso al diccionario completo.
    """
    # Aplicar estilo de párrafo
    if 'estilo_parrafo' in merged_config:
        apply_cell_paragraph_style(cell, doc, merged_config['estilo_parrafo'])
    
    # Aplicar color de fondo
    if 'color_fondo' in merged_config:
        apply_cell_background_color(cell, merged_config['color_fondo'])
    
    # Aplicar alineación vertical
    # IMPORTANTE: Se aplica DESPUÉS del estilo de párrafo para evitar el bug
    # documentado donde aplicar estilos "daña" la alineación vertical
    if 'alineacion_vertical' in merged_config:
        apply_cell_vertical_alignment(cell, merged_config['alineacion_vertical'])
    
    # Aplicar alineación horizontal
    if 'alineacion_horizontal' in merged_config:
        apply_cell_horizontal_alignment(cell, merged_config['alineacion_horizontal'])
    
    # Aplicar color de texto
    if 'color_texto' in merged_config:
        apply_cell_text_color(cell, merged_config['color_texto'])
