"""
Internal helper functions for figure manipulation in Word templates.

This module contains low-level functions for inserting figures with/without captions,
creating figure captions with SEQ fields, bookmarks, and cross-references.

Private module - Not intended for direct external use.
Import from the public API in api.py instead.
"""

from genericpath import exists
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


def aux_insertar_figura_sin_titulo(paragraph, key, lista_figuras):
    """Inserta imágenes en un párrafo de Word SIN título/Caption.
    
    Args:
        paragraph: El párrafo del documento Word donde se buscará el marcador.
        key: El marcador de posición a buscar (por ejemplo, "<<fig_ejecucion_campania>>").
        lista_figuras: Lista de diccionarios con keys "ruta", "titulo", "tamanio", "bookmark".
                      Para figuras sin título, titulo y bookmark deben ser "".
    
    Returns:
        True si se insertaron las imágenes, False si no se encontró el marcador.
        
    Ejemplo de uso:
        lista_figuras = [
            {"ruta": "img1.png", "titulo": "", "tamanio": 6, "bookmark": ""},
            {"ruta": "img2.png", "titulo": "", "tamanio": 5, "bookmark": ""}
        ]
    """
    full_text = "".join(run.text for run in paragraph.runs)
    
    if key not in full_text:
        return False
    
    # Limpiar el párrafo (borrar todos los runs)
    for run in paragraph.runs:
        run.text = ""
    
    # Obtener referencias para inserción
    p_element = paragraph._element
    parent = p_element.getparent()
    indice_base = parent.index(p_element)
    
    # Insertar todas las figuras sin título
    offset = 0
    for idx, item in enumerate(lista_figuras):
        ruta = item.get("ruta", "")
        ancho = item.get("tamanio", 6)
        
        # Si es la primera imagen, usar el párrafo actual
        if offset == 0:
            # Aplicar estilo "Figura" al párrafo actual
            p_element = paragraph._element
            pPr = p_element.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p_element.insert(0, pPr)
            # Aplicar estilo Figura
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is None:
                pStyle = OxmlElement('w:pStyle')
                pPr.insert(0, pStyle)
            pStyle.set(qn('w:val'), 'Figura')
            
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            if exists(ruta):
                run.add_picture(ruta, width=Inches(ancho))
            offset = 1
        else:
            # Crear nuevo párrafo para imagen con estilo "Figura"
            nuevo_p_img = OxmlElement('w:p')
            # Propiedades del párrafo
            pPr_img = OxmlElement('w:pPr')
            pStyle_img = OxmlElement('w:pStyle')
            pStyle_img.set(qn('w:val'), 'Figura')
            pPr_img.append(pStyle_img)
            nuevo_p_img.append(pPr_img)
            
            parent.insert(indice_base + offset, nuevo_p_img)
            # Convertir a Paragraph para poder agregar imagen
            para_img = Paragraph(nuevo_p_img, paragraph._parent)
            if exists(ruta):
                para_img.add_run().add_picture(ruta, width=Inches(ancho))
            offset += 1
    
    return True


def aux_insertar_figuras_con_titulo(paragraph, key, lista_figuras):
    """Inserta MÚLTIPLES imágenes con pies de figura válidos (Caption + SEQ + Bookmarks).
    
    Args:
        paragraph: El párrafo del documento Word donde se buscará el marcador.
        key: El marcador de posición a buscar (por ejemplo, "<<fig_ejecucion_campania>>").
        lista_figuras: Lista de diccionarios con keys "ruta", "titulo", "tamanio", "bookmark" (opcional).
    
    Returns:
        Lista de nombres de bookmarks creados si se insertaron las imágenes, False si no se encontró el marcador.
        
    Ejemplo de uso:
        lista_figuras = [
            {"ruta": "img1.png", "titulo": "Mapa de ubicación", "tamanio": 6, "bookmark": "_Ref_Mapa1"},
            {"ruta": "img2.png", "titulo": "Temperatura del agua", "tamanio": 5}
        ]
    """
    full_text = "".join(run.text for run in paragraph.runs)
    
    if key not in full_text:
        return False
    
    # Limpiar el párrafo (borrar todos los runs)
    for run in paragraph.runs:
        run.text = ""
    
    # Obtener referencias para inserción
    p_element = paragraph._element
    parent = p_element.getparent()
    indice_base = parent.index(p_element)
    
    # Lista para almacenar los bookmarks creados
    bookmarks_creados = []
    
    # Insertar todas las figuras con sus títulos
    offset = 0
    for idx, item in enumerate(lista_figuras):
        ruta = item.get("ruta", "")
        titulo = item.get("titulo", f"Figura {idx + 1}")
        ancho = item.get("tamanio", 6)
        bookmark = item.get("bookmark", None)  # Bookmark personalizado opcional
        
        # Si es la primera imagen, usar el párrafo actual
        if offset == 0:
            # Aplicar estilo "Figura" al párrafo actual
            p_element = paragraph._element
            pPr = p_element.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p_element.insert(0, pPr)
            # Aplicar estilo Figura
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is None:
                pStyle = OxmlElement('w:pStyle')
                pPr.insert(0, pStyle)
            pStyle.set(qn('w:val'), 'Figura')
            
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            if exists(ruta):
                run.add_picture(ruta, width=Inches(ancho))
            offset = 1
        else:
            # Crear nuevo párrafo para imagen con estilo "Figura"
            nuevo_p_img = OxmlElement('w:p')
            # Propiedades del párrafo
            pPr_img = OxmlElement('w:pPr')
            pStyle_img = OxmlElement('w:pStyle')
            pStyle_img.set(qn('w:val'), 'Figura')
            pPr_img.append(pStyle_img)
            nuevo_p_img.append(pPr_img)
            
            parent.insert(indice_base + offset, nuevo_p_img)
            # Convertir a Paragraph para poder agregar imagen
            para_img = Paragraph(nuevo_p_img, paragraph._parent)
            if exists(ruta):
                para_img.add_run().add_picture(ruta, width=Inches(ancho))
            offset += 1
        
        # Usar crear_pie_de_figura con bookmark
        nuevo_p_caption, bookmark_name = crear_pie_de_figura(
            parent, 
            indice_base + offset, 
            titulo,
            bookmark_name=bookmark
        )
        bookmarks_creados.append(bookmark_name)
        offset += 1
        
        # Salto de línea
        nuevo_p_salto = OxmlElement('w:p')
        parent.insert(indice_base + offset, nuevo_p_salto)
        offset += 1
    
    return bookmarks_creados


def aux_insertar_referencia_cruzada(paragraph, nombre_bookmark, texto_antes="Figura", mostrar_numero=True):
    """Inserta una referencia cruzada a una figura en un párrafo.
    
    Args:
        paragraph: El párrafo donde insertar la referencia (objeto Paragraph)
        nombre_bookmark: Nombre del bookmark de la figura a referenciar
        texto_antes: Texto a mostrar antes del número (default: "Figura")
        mostrar_numero: Si True, muestra el número de la figura; si False, solo el bookmark
    
    Ejemplo de uso:
        # Para: "Como se muestra en la Figura 3"
        p = doc.add_paragraph("Como se muestra en la ")
        aux_insertar_referencia_cruzada(p, "_Ref_Fig_Mi_Figura", "Figura")
    """
    # Agregar texto antes si se proporciona
    if texto_antes:
        run_texto = paragraph.add_run(texto_antes + " ")
    
    # Crear el run para el campo REF
    run = paragraph.add_run()
    
    # Crear el campo REF para la figura
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    # \r muestra solo el número de secuencia (SEQ), \h hace el campo clickable (hyperlink)
    if mostrar_numero:
        instrText.text = f' REF {nombre_bookmark} \\h '
    else:
        instrText.text = f' REF {nombre_bookmark} \\r \\h '
    
    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    
    # Texto por defecto (se actualizará en Word con F9)
    text_run = OxmlElement('w:r')
    text_elem = OxmlElement('w:t')
    text_elem.text = "XX"  # Placeholder que Word actualizará
    text_run.append(text_elem)
    
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    
    # Agregar elementos al run
    run._element.append(fldChar_begin)
    run._element.append(instrText)
    run._element.append(fldChar_separate)
    run._element.append(text_run)
    run._element.append(fldChar_end)


def crear_pie_de_figura(parent, indice, titulo, bookmark_name=None):
    """Crea un párrafo con pie de figura válido para Word con estilo, campo SEQ y bookmark.
    
    Args:
        parent: El elemento padre XML donde insertar
        indice: Posición donde insertar
        titulo: Texto descriptivo de la figura
        bookmark_name: Nombre del bookmark para referencias cruzadas (opcional, se genera automático si None)
    
    Returns:
        Tupla (elemento XML del párrafo creado, nombre del bookmark)
    
    Estilos aplicados:
        - Si longitud total < 115 caracteres: usa estilo "Car_centrado" (style_id: 'Carcentrado')
        - Si longitud total >= 115 caracteres: usa estilo "Car_justificado" (style_id: 'Carjustificado')
        - Longitud total = "Figura XXX. " + titulo (estimado ~13 + len(titulo))
    
    Bookmark:
        El bookmark se crea alrededor del NÚMERO únicamente, no incluye "Figura" ni el título.
        Estructura: "Figura " [bookmark_start] "8" [bookmark_end] ". Título"
        Esto permite que las referencias cruzadas muestren solo el número.
    """
    cantidad_de_caracteres = 115  # Umbral para decidir entre centrado o justificado
    
    # Generar nombre de bookmark si no se proporciona
    if bookmark_name is None:
        # Usar los primeros 30 caracteres del título, reemplazando espacios y caracteres especiales
        bookmark_name = f"_Ref_Fig_{titulo[:30].replace(' ', '_').replace(',', '').replace('.', '')}"
    
    # Generar ID único para el bookmark basado en el hash del nombre
    bookmark_id = str(abs(hash(bookmark_name)) % 1000000)
    
    # Crear párrafo con estilo
    nuevo_p = OxmlElement('w:p')
    
    # Determinar qué estilo usar según la longitud del título
    longitud_estimada = 13 + len(titulo)
    estilo_titulo = 'Carcentrado' if longitud_estimada < cantidad_de_caracteres else 'Carjustificado'
    
    # Propiedades del párrafo
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), estilo_titulo)
    pPr.append(pStyle)
    nuevo_p.append(pPr)
    
    # Run para "Figura " (fuera del bookmark)
    run1 = OxmlElement('w:r')
    text1 = OxmlElement('w:t')
    text1.set(qn('xml:space'), 'preserve')
    text1.text = 'Figura '
    run1.append(text1)
    nuevo_p.append(run1)
    
    # Bookmark start (antes del número, después de "Figura ")
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), bookmark_id)
    bookmark_start.set(qn('w:name'), bookmark_name)
    nuevo_p.append(bookmark_start)
    
    # Campo SEQ para numeración automática
    run_begin = OxmlElement('w:r')
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run_begin.append(fldChar_begin)
    nuevo_p.append(run_begin)
    
    run_instr = OxmlElement('w:r')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' SEQ Figura \\* ARABIC '
    run_instr.append(instrText)
    nuevo_p.append(run_instr)
    
    run_separate = OxmlElement('w:r')
    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    run_separate.append(fldChar_separate)
    nuevo_p.append(run_separate)
    
    run_result = OxmlElement('w:r')
    text_result = OxmlElement('w:t')
    text_result.text = '1'  # Placeholder que Word actualizará
    run_result.append(text_result)
    nuevo_p.append(run_result)
    
    run_end = OxmlElement('w:r')
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run_end.append(fldChar_end)
    nuevo_p.append(run_end)
    
    # Bookmark end (después del número, antes del título)
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), bookmark_id)
    nuevo_p.append(bookmark_end)
    
    # Run para el texto descriptivo (fuera del bookmark)
    run2 = OxmlElement('w:r')
    text2 = OxmlElement('w:t')
    text2.set(qn('xml:space'), 'preserve')
    text2.text = f'. {titulo}'
    run2.append(text2)
    nuevo_p.append(run2)
    
    # Insertar en el documento
    parent.insert(indice, nuevo_p)
    
    return nuevo_p, bookmark_name
