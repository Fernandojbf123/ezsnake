"""
Public API for Word Template Writer Module
==========================================

This module provides high-level orchestrator functions (in Spanish) for manipulating
Word templates. These are the main functions users should call.

Functions:
    - reemplazar_variable_por_figura: Insert figures with/without captions
    - reemplazar_referencias_cruzadas_de_figuras: Create cross-references to figures
    - reemplazar_variable_por_tabla: Fill table placeholders and prepare table references
    - reemplazar_referencias_cruzadas_de_tablas: Create cross-references to tables
    - reemplazar_texto_en_plantilla: Replace text variables in template
    - insertar_lista_en_plantilla: Insert bulleted lists in template
    - insertar_documento_externo_en_plantilla: Insert external Word documents
    - rellenar_tablas_en_plantilla: Fill tables with DataFrame data
"""

# Import helper functions from private modules
from ._figure_helpers import (
    aux_insertar_figura_sin_titulo,
    aux_insertar_figuras_con_titulo,
    aux_insertar_referencia_cruzada,
)
from ._text_helpers import replace_text_variables_in_paragraph, replace_text_variables_in_tables
from ._document_helpers import insert_external_document
from ._table_helpers import fill_table, insertar_titulo_de_tabla_con_bookmark


def reemplazar_variable_por_figura(doc, diccionario_de_reemplazos: dict):
    """Inserta todas las figuras definidas en el diccionario en la plantilla de Word.
    
    Muta el diccionario de entrada al unirlo con la información de los bookmarks 
    creados para referencias cruzadas.
        
    Args:
        doc: Objeto Document de python-docx
        diccionario_de_reemplazos: Diccionario donde las claves que comienzan con "<<fig" 
                                   contienen listas de figuras a insertar.
    
    Returns:
        None (el diccionario de entrada se muta in-place agregando keys "<<reffigura*>>")
    
    Casos soportados:
        - Caso 1: Figuras SIN título - titulo="" y bookmark=""
                 Se insertan imágenes sin Caption, numeración ni bookmarks
        - Caso 2: Figuras CON título - titulo y bookmark con contenido
                 Se crean pies de figura con estilo Caption, campo SEQ y bookmarks
    
    Estructura esperada del diccionario:
        {
            # Caso 1: Figuras SIN título
            "<<fig_fotos>>": [
                {"ruta": "foto1.png", "titulo": "", "tamanio": 6, "bookmark": "", "estilo_figura": "Figura", "estilo_titulo": "Normal"},
                {"ruta": "foto2.png", "titulo": "", "tamanio": 5, "bookmark": "", "estilo_figura": "Figura", "estilo_titulo": "Normal"}
            ],
            
            # Caso 2: Figuras CON título
            "<<fig_mapas>>": [
                {"ruta": "mapa1.png", "titulo": "Mapa de ubicación", "tamanio": 6, "bookmark": "RefFigura_Mapa1", "estilo_figura": "Figura", "estilo_titulo": "Normal"},
                {"ruta": "mapa2.png", "titulo": "Temperatura del agua", "tamanio": 5, "bookmark": "RefFigura_Temp", "estilo_figura": "Figura", "estilo_titulo": "Normal"}
            ],
            
            "<<orden_servicio>>": "12345",  # Variables no-figura se ignoran aquí
        }
    
    Ejemplo de uso:
        from docx import Document
        from word_template_writer import reemplazar_variable_por_figura, reemplazar_referencias_cruzadas_de_figuras
        
        doc = Document('plantilla.docx')
        diccionario = {
            "<<fig_mapas>>": [
                {"ruta": "mapa1.png", "titulo": "Ubicación sondas", "tamanio": 6, "bookmark": "RefFigura_Mapa1", "estilo_figura": "Figura", "estilo_titulo": "Normal"},
                {"ruta": "mapa2.png", "titulo": "Temperatura", "tamanio": 5, "bookmark": "RefFigura_Temp", "estilo_figura": "Figura", "estilo_titulo": "Normal"}
            ],
            "<<fig_fotos>>": [
                {"ruta": "foto1.png", "titulo": "", "tamanio": 4, "bookmark": "", "estilo_figura": "Figura", "estilo_titulo": "Normal"}
            ]
        }
        
        # Paso 1: Insertar las figuras (agrega keys <<reffigura*>> al diccionario)
        reemplazar_variable_por_figura(doc, diccionario)
        # El diccionario ahora contiene:
        # {
        #     "<<reffigura_mapas>>": ["RefFigura_Mapa1", "RefFigura_Temp"],
        #     "<<reffigura_fotos>>": None,
        #     ... (keys originales se mantienen)
        # }
        
        # Paso 2: Insertar referencias cruzadas (si hay marcadores <<reffigura*>> en la plantilla)
        # Ejemplo: "De la <<reffigura_mapas>> se observa..." → "De la Figura 1 a la 2 se observa..."
        reemplazar_referencias_cruzadas_de_figuras(doc, diccionario)
        
        doc.save('documento_con_figuras.docx')
    """
    bookmarks_info = {}
    
    if diccionario_de_reemplazos is None:
        raise ValueError("El diccionario de reemplazos no puede ser None.")
    
    # Filtrar solo las variables que son figuras (comienzan con "<<fig")
    variables_figuras = {k: v for k, v in diccionario_de_reemplazos.items() if k.startswith("<<fig")}
    
    # Procesar cada variable de figura
    for variable, datos_figuras in variables_figuras.items():
        if not isinstance(datos_figuras, list):
            print(f"Advertencia: La variable '{variable}' no contiene una lista. Se omite.")
            continue
        
        if len(datos_figuras) == 0:
            print(f"Advertencia: La variable '{variable}' contiene una lista vacía. Se omite.")
            continue
        
        # Determinar si son figuras CON o SIN título
        # Caso 1: Figuras SIN título - todos los títulos están vacíos
        tiene_titulo = any(item.get("titulo", "") != "" for item in datos_figuras)
        
        # Buscar el marcador en todos los párrafos del documento
        for parrafo in doc.paragraphs:
            if variable in parrafo.text:
                
                # Crear key de referencia: "<<fig_mapas>>" -> "<<reffigura_mapas>>"
                variable_ref_key = variable.replace("<<fig", "<<reffigura", 1)
                
                if not tiene_titulo:
                    # Caso 1: Figuras SIN título (no Caption, no bookmark)
                    resultado = aux_insertar_figura_sin_titulo(parrafo, variable, datos_figuras)
                    if resultado:
                        bookmarks_info[variable_ref_key] = None  # Sin bookmarks para figuras sin título
                        break  # Ya se insertó, pasar a la siguiente variable
                else:
                    # Caso 2: Figuras CON título (Caption + SEQ + Bookmarks)
                    bookmarks_creados = aux_insertar_figuras_con_titulo(parrafo, variable, datos_figuras)
                    if bookmarks_creados:
                        bookmarks_info[variable_ref_key] = bookmarks_creados
                        break  # Ya se insertó, pasar a la siguiente variable
    
    # Mutar el diccionario agregando información de bookmarks
    diccionario_de_reemplazos.update(bookmarks_info)
    
    msg = "Figuras insertadas y bookmarks creados para referencias cruzadas"
    print(msg)


def reemplazar_referencias_cruzadas_de_figuras(doc, diccionario_de_reemplazos: dict):
    """Reemplaza marcadores <<reffigura*>> con referencias cruzadas a figuras.
    
    Args:
        doc: Objeto Document de python-docx
        diccionario_de_reemplazos: Diccionario retornado/mutado por reemplazar_variable_por_figura
                      Formato: {"<<reffigura_demo>>": ["RefFigura_Mapa1", "RefFigura_Temp"], ...}
    
    Comportamiento:
        - Busca variables que empiecen con "<<reffigura" en los párrafos
        - Si la lista tiene 1 bookmark: inserta "Figura X"
        - Si la lista tiene 2+ bookmarks: inserta "Figura X a la Y"
        - X e Y son referencias cruzadas reales (campos REF) que muestran solo el número
        - Procesa TODOS los marcadores de un párrafo en una sola pasada
    
    Nota importante:
        Los bookmarks creados por crear_pie_de_figura solo incluyen el número de la figura,
        no el texto "Figura" ni el título. Por eso las referencias muestran solo el número.
    
    Ejemplo:
        Entrada en plantilla: "De la <<reffigura_demo_con_titulo>> se muestra el poder."
        Salida: "De la Figura 8 a la 10 se muestra el poder."
                (donde "8" y "10" son campos REF clickeables que muestran solo el número)
    """
    # Filtrar solo variables con bookmarks válidos
    variables_validas = {k: v for k, v in diccionario_de_reemplazos.items() 
                        if v is not None and k.startswith("<<reffigura")}
    
    # Para cada párrafo
    for parrafo in doc.paragraphs:
        full_text = "".join(run.text for run in parrafo.runs)
        
        # Encontrar TODOS los marcadores <<reffigura*>> en este párrafo
        marcadores_en_parrafo = []
        for variable_ref, lista_bookmarks in variables_validas.items():
            if variable_ref in full_text:
                # Encontrar todas las ocurrencias del marcador en el párrafo
                pos = full_text.find(variable_ref)
                if pos != -1:
                    marcadores_en_parrafo.append((pos, variable_ref, lista_bookmarks))
        
        # Si no hay marcadores en este párrafo, continuar al siguiente
        if not marcadores_en_parrafo:
            continue
        
        # Ordenar marcadores por posición (de izquierda a derecha)
        marcadores_en_parrafo.sort(key=lambda x: x[0])
        
        # Limpiar todos los runs del párrafo
        for run in parrafo.runs:
            run.text = ""
        
        # Reconstruir el párrafo procesando todos los marcadores
        pos_actual = 0
        
        for pos_marcador, variable_ref, lista_bookmarks in marcadores_en_parrafo:
            # Agregar texto antes del marcador
            if pos_marcador > pos_actual:
                texto_antes = full_text[pos_actual:pos_marcador]
                parrafo.add_run(texto_antes)
            
            # Insertar la referencia cruzada
            primer_bookmark = lista_bookmarks[0]
            ultimo_bookmark = lista_bookmarks[-1]
            
            if len(lista_bookmarks) == 1:
                # Caso: Solo una figura - "Figura X"
                aux_insertar_referencia_cruzada(parrafo, primer_bookmark, texto_antes="Figura", mostrar_numero=True)
            else:
                # Caso: Múltiples figuras - "Figura X a la Y"
                aux_insertar_referencia_cruzada(parrafo, primer_bookmark, texto_antes="Figura", mostrar_numero=True)
                parrafo.add_run(" a la ")
                aux_insertar_referencia_cruzada(parrafo, ultimo_bookmark, texto_antes="", mostrar_numero=True)
            
            # Avanzar posición actual
            pos_actual = pos_marcador + len(variable_ref)
        
        # Agregar texto después del último marcador
        if pos_actual < len(full_text):
            texto_despues = full_text[pos_actual:]
            parrafo.add_run(texto_despues)

    msg = "Referencias cruzadas insertadas."
    print(msg)


def reemplazar_referencias_cruzadas_de_tablas(doc, diccionario_de_reemplazos: dict):
    """Reemplaza marcadores <<reftabla_*>> con referencias cruzadas a tablas.

    Args:
        doc: Objeto Document de python-docx
        diccionario_de_reemplazos: Diccionario mutado por reemplazar_variable_por_tabla
                                  Formato: {"<<reftabla_demo>>": ["Reftabla_MiTabla"], ...}

    Comportamiento:
        - Busca variables que empiecen con "<<reftabla_" en los párrafos
        - Si la lista tiene 1 bookmark: inserta "Tabla X"
        - Si la lista tiene 2+ bookmarks: inserta "Tabla X a la Y"
        - X e Y son referencias cruzadas reales (campos REF) que muestran el número/etiqueta
        - Procesa TODOS los marcadores de un párrafo en una sola pasada

    Ejemplo de uso:
        from docx import Document
        from word_template_writer import reemplazar_referencias_cruzadas_de_tablas

        doc = Document('plantilla.docx')
        diccionario = {
            "<<reftabla_resultados>>": ["Reftabla_Resultados_1"],
            "<<reftabla_series>>": ["Reftabla_Serie_A", "Reftabla_Serie_B"],
        }

        reemplazar_referencias_cruzadas_de_tablas(doc, diccionario)
        doc.save('documento_con_refs_tablas.docx')
    """
    variables_validas = {
        k: v for k, v in diccionario_de_reemplazos.items()
        if v is not None and k.startswith("<<reftabla_")
    }

    for parrafo in doc.paragraphs:
        full_text = "".join(run.text for run in parrafo.runs)

        marcadores_en_parrafo = []
        for variable_ref, lista_bookmarks in variables_validas.items():
            if variable_ref in full_text:
                pos = full_text.find(variable_ref)
                if pos != -1:
                    marcadores_en_parrafo.append((pos, variable_ref, lista_bookmarks))

        if not marcadores_en_parrafo:
            continue

        marcadores_en_parrafo.sort(key=lambda x: x[0])

        for run in parrafo.runs:
            run.text = ""

        pos_actual = 0

        for pos_marcador, variable_ref, lista_bookmarks in marcadores_en_parrafo:
            if pos_marcador > pos_actual:
                texto_antes = full_text[pos_actual:pos_marcador]
                parrafo.add_run(texto_antes)

            primer_bookmark = lista_bookmarks[0]
            ultimo_bookmark = lista_bookmarks[-1]

            if len(lista_bookmarks) == 1:
                aux_insertar_referencia_cruzada(parrafo, primer_bookmark, texto_antes="Tabla", mostrar_numero=True)
            else:
                aux_insertar_referencia_cruzada(parrafo, primer_bookmark, texto_antes="Tabla", mostrar_numero=True)
                parrafo.add_run(" a la ")
                aux_insertar_referencia_cruzada(parrafo, ultimo_bookmark, texto_antes="", mostrar_numero=True)

            pos_actual = pos_marcador + len(variable_ref)

        if pos_actual < len(full_text):
            texto_despues = full_text[pos_actual:]
            parrafo.add_run(texto_despues)

    msg = "Referencias cruzadas de tablas insertadas."
    print(msg)


def reemplazar_variable_por_tabla(doc, diccionario_de_reemplazos: dict):
    """Reemplaza variables <<nuevatabla_*>> por tablas y prepara referencias <<reftabla_*>>.

    Args:
        doc: Objeto Document de python-docx
        diccionario_de_reemplazos: Diccionario con variables del documento

    Estructura esperada de cada variable de tabla:
        {
            "tabla": pd.DataFrame,
            "estilos_de_tabla": dict,
            "titulo": "Texto del título de tabla",
            "bookmark": "Reftabla_MiTabla"
        }

    Comportamiento:
        - Busca keys que comienzan con "<<nuevatabla_"
        - Inserta título con bookmark antes de la tabla objetivo
        - Rellena la tabla con fill_table()
        - Muta diccionario_de_reemplazos agregando "<<reftabla_*>>": [bookmark]

    Ejemplo de uso:
        import pandas as pd
        from docx import Document
        from word_template_writer import EstilosTabla, reemplazar_variable_por_tabla

        doc = Document('plantilla.docx')
        df = pd.DataFrame({"A": [1, 2], "B": [3, 4]})
        estilos = EstilosTabla(doc).to_dict()

        diccionario = {
            "<<nuevatabla_resultados>>": {
                "tabla": df,
                "estilos_de_tabla": estilos,
                "titulo": "Tabla 1. Resultados del análisis",
                "bookmark": "Reftabla_Resultados_1",
            }
        }

        reemplazar_variable_por_tabla(doc, diccionario)
        doc.save('documento_con_tablas.docx')
    """
    if diccionario_de_reemplazos is None:
        raise ValueError("El diccionario de reemplazos no puede ser None.")

    variables_tabla = {
        k: v for k, v in diccionario_de_reemplazos.items()
        if k.startswith("<<nuevatabla_")
    }

    referencias_tablas = {}

    for variable, config_tabla in variables_tabla.items():

        if not isinstance(config_tabla, dict):
            raise ValueError(
                f"La variable '{variable}' debe contener un diccionario con configuración de tabla."
            )

        keys_requeridas = ["tabla", "estilos_de_tabla", "titulo", "bookmark"]
        faltantes = [key for key in keys_requeridas if key not in config_tabla]
        if faltantes:
            raise ValueError(
                f"La variable '{variable}' no tiene las keys requeridas: {faltantes}."
            )

        tabla_df = config_tabla["tabla"]
        estilos_de_tabla = config_tabla["estilos_de_tabla"]
        titulo_tabla = config_tabla["titulo"]
        bookmark_tabla = config_tabla["bookmark"]

        if not isinstance(bookmark_tabla, str) or not bookmark_tabla.startswith("Reftabla"):
            raise ValueError(
                f"El bookmark de '{variable}' debe comenzar con 'Reftabla'."
            )

        insertar_titulo_de_tabla_con_bookmark(
            doc,
            marcador_tabla=variable,
            titulo=titulo_tabla,
            bookmark=bookmark_tabla,
            estilo_titulo="Normal",
        )

        fill_table(
            doc,
            variable,
            tabla_df,
            estilos_de_tabla,
            None,
        )

        variable_ref_key = variable.replace("<<nuevatabla_", "<<reftabla_", 1)
        referencias_tablas[variable_ref_key] = [bookmark_tabla]

    diccionario_de_reemplazos.update(referencias_tablas)

    msg = "Tablas insertadas y referencias de tabla preparadas en el diccionario."
    print(msg)


def reemplazar_texto_en_plantilla(doc, diccionario_de_reemplazos):
    """Reemplaza los marcadores de posición de texto en un documento de Word.
    
    Args:
        doc: El documento de Word (objeto Document de python-docx).
        diccionario_de_reemplazos: Diccionario donde las claves son los marcadores de posición 
                                  a buscar (por ejemplo, "<<orden_de_servicio>>") y los valores 
                                  son los textos que los reemplazarán (ej: "12345").
                                  Los valores pueden ser strings o listas de strings.
    
        Comportamiento:
                - Ignora marcadores reservados para:
                    * Figuras (<<fig...>>)
                    * Referencias de figuras (<<reffigura...>>)
                    * Tablas (<<nuevatabla_...>>, <<editartabla...>>)
                    * Referencias de tablas (<<reftabla_...>>)
                    * Documentos externos (<<external_doc_...>>)
        - Procesa todas las variables de texto en cada párrafo de una sola vez
        - Preserva el formato del primer run del párrafo
        - Si el valor es una lista y el marcador ocupa todo el párrafo:
                    * Se crean múltiples párrafos (uno por cada elemento de la lista)
                    * El valor debe ser una lista de tuplas: (texto, estilo)
                    * Si estilo == "", se aplica estilo "Normal"
                    * Si texto == "", se inserta un salto de línea (párrafo vacío)
        - Si hay texto adicional o múltiples marcadores en el mismo párrafo:
          * Solo se usa el primer elemento de la lista
    
    Ejemplo de uso:
        from docx import Document
        from word_template_writer import reemplazar_texto_en_plantilla
        
        doc = Document('plantilla.docx')
        
        # Caso 1: Reemplazo simple con strings
        diccionario = {
            "<<orden_servicio>>": "12345",
            "<<cliente>>": "ACME Corporation",
            "<<fecha>>": "21/05/2026",
        }
        
        # Caso 2: Reemplazo con múltiples párrafos (lista de tuplas)
        diccionario = {
            "<<seccion_1>>": [
                ("este es el primer párrafo", "estilo 2"),
                ("este es el segundo párrafo", ""),
                ("", ""),
                ("tercer párrafo", "estilo 3"),
            ],
            "<<orden_servicio>>": "12345",
        }
        
        reemplazar_texto_en_plantilla(doc, diccionario)
        doc.save('documento_con_texto.docx')
    
    Nota:
        Para que se creen múltiples párrafos, el marcador debe ser el único
        contenido del párrafo en la plantilla (sin texto adicional).
    """
    # Filtrar solo variables de texto, excluyendo prefijos reservados de otras operaciones.
    prefijos_reservados = (
        "<<fig",
        "<<reffigura",
        "<<nuevatabla_",
        "<<editartabla",
        "<<reftabla_",
        "<<external_doc_",
        "<<lista_",
        "<<ul_",
    )
    variables_texto = {
        k: v for k, v in diccionario_de_reemplazos.items()
        if not any(k.startswith(prefijo) for prefijo in prefijos_reservados)
    }
    
    # Para cada párrafo, procesar TODAS las variables de texto de una sola vez
    for parrafo in doc.paragraphs:
        # Encontrar todas las variables que están en este párrafo
        variables_en_parrafo = []
        for variable, dato in variables_texto.items():
            if variable in parrafo.text:
                variables_en_parrafo.append((variable, dato))
        
        # Si hay variables en este párrafo, reemplazarlas todas de una vez
        if variables_en_parrafo:
            replace_text_variables_in_paragraph(parrafo, variables_en_parrafo)
    
    msg = "Se agregaron los textos al documento."
    print(msg)


def insertar_lista_en_plantilla(doc, diccionario_de_reemplazos, nombre_estilo="List Bullet"):
    """Inserta listas con viñetas en la plantilla de Word.
    
    Args:
        doc: El documento de Word (objeto Document de python-docx).
        diccionario_de_reemplazos: Diccionario donde las claves que comienzan con "<<lista_" 
                                   contienen listas de elementos a insertar como viñetas.
        nombre_estilo: Nombre del estilo de párrafo a aplicar a cada elemento de la lista.
                      Por defecto usa "List Bullet" (estilo incorporado de Word).
                      Se recomienda crear un estilo personalizado llamado "unsorted_list" 
                      en la plantilla para mayor control del formato.
    
    Comportamiento:
        - Busca marcadores que empiecen con "<<lista_" en los párrafos
        - El marcador DEBE estar solo en el párrafo (sin texto adicional)
        - Si el marcador no está solo, lanza ValueError
        - Reemplaza el párrafo completo con múltiples párrafos (uno por elemento)
        - Aplica el estilo especificado a cada párrafo nuevo
        - Preserva la posición original del marcador en el documento
    
    Estructura esperada del diccionario:
        {
            "<<lista_resultados>>": [
                "Primer resultado del análisis",
                "Segundo resultado encontrado",
                "Tercer resultado importante"
            ],
            "<<lista_recomendaciones>>": [
                "Primera recomendación",
                "Segunda recomendación"
            ],
            "<<orden_servicio>>": "12345"  # Variables no-lista se ignoran aquí
        }
    
    Ejemplo de uso:
        from docx import Document
        from word_template_writer import insertar_lista_en_plantilla
        
        doc = Document('plantilla.docx')
        diccionario = {
            "<<lista_hallazgos>>": [
                "Se observó un incremento del 15% en la temperatura",
                "La salinidad mostró valores consistentes con el promedio histórico",
                "Se detectaron concentraciones elevadas de clorofila-a"
            ],
            "<<lista_equipos>>": [
                "CTD SBE 911plus",
                "ADCP Workhorse 300 kHz",
                "Fluorómetro WETLabs ECO"
            ]
        }
        
        # Usando estilo por defecto
        insertar_lista_en_plantilla(doc, diccionario)
        
        # Usando estilo personalizado (debe existir en la plantilla)
        insertar_lista_en_plantilla(doc, diccionario, nombre_estilo="unsorted_list")
        
        doc.save('documento_con_listas.docx')
    
    Raises:
        ValueError: Si el marcador no está solo en el párrafo o si el valor no es una lista
    
    Nota importante:
        Para que esta función trabaje correctamente, el marcador <<lista_*>> debe ser
        el ÚNICO contenido del párrafo en la plantilla. Por ejemplo:
        
        ✓ CORRECTO:   Párrafo que solo contiene "<<lista_hallazgos>>"
        ✗ INCORRECTO: "Los hallazgos son: <<lista_hallazgos>>"
    """
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph
    
    # Filtrar solo variables que son listas (comienzan con "<<lista_")
    variables_listas = {k: v for k, v in diccionario_de_reemplazos.items() if k.startswith("<<lista_")}
    
    # Procesar cada variable de lista
    for variable, elementos_lista in variables_listas.items():
        # Validar que el valor sea una lista
        if not isinstance(elementos_lista, list):
            raise ValueError(
                f"El valor de '{variable}' debe ser una lista. "
                f"Se recibió: {type(elementos_lista).__name__}"
            )
        
        if len(elementos_lista) == 0:
            print(f"Advertencia: La variable '{variable}' contiene una lista vacía. Se omite.")
            continue
        
        # Buscar el marcador en todos los párrafos del documento
        for parrafo in doc.paragraphs:
            if variable in parrafo.text:
                # Obtener el texto completo del párrafo
                full_text = "".join(run.text for run in parrafo.runs).strip()
                
                # Validar que el marcador esté SOLO en el párrafo
                if full_text != variable:
                    raise ValueError(
                        f"El marcador '{variable}' debe estar solo en el párrafo. "
                        f"Texto encontrado: '{full_text}'. "
                        f"Los elementos <<lista_*>> deben estar solos en un párrafo."
                    )
                
                # Reemplazar el primer párrafo con el primer elemento de la lista
                primer_run = parrafo.runs[0] if parrafo.runs else parrafo.add_run()
                primer_run.text = str(elementos_lista[0])
                
                # Limpiar el resto de runs del párrafo
                for i in range(1, len(parrafo.runs)):
                    parrafo.runs[i].text = ""
                
                # Aplicar el estilo al primer párrafo
                try:
                    parrafo.style = nombre_estilo
                except KeyError:
                    print(f"Advertencia: El estilo '{nombre_estilo}' no existe en el documento. "
                          f"Se mantiene el estilo original del párrafo.")
                
                # Obtener el elemento actual y su padre
                elemento_actual = parrafo._element
                
                # Insertar párrafos adicionales para el resto de elementos
                for texto in elementos_lista[1:]:
                    # Crear un nuevo elemento de párrafo
                    nuevo_elemento = OxmlElement('w:p')
                    
                    # Insertar el nuevo párrafo después del anterior
                    elemento_actual.addnext(nuevo_elemento)
                    
                    # Crear objeto Paragraph desde el elemento XML
                    nuevo_parrafo = Paragraph(nuevo_elemento, parrafo._parent)
                    
                    # Aplicar el estilo
                    try:
                        nuevo_parrafo.style = nombre_estilo
                    except KeyError:
                        pass  # Si falla, usa el estilo por defecto
                    
                    # Agregar el texto al nuevo párrafo
                    nuevo_parrafo.add_run(str(texto))
                    
                    # Actualizar el elemento actual para la siguiente iteración
                    elemento_actual = nuevo_elemento
                
                # Ya se procesó este marcador, pasar al siguiente
                break
    
    msg = "Listas insertadas en el documento."
    print(msg)


def insertar_documento_externo_en_plantilla(doc, diccionario_de_reemplazos):
    """Inserta uno o más documentos Word externos en la plantilla.
    
    Args:
        doc: El documento de Word (objeto Document de python-docx).
        diccionario_de_reemplazos: Diccionario que debe contener keys que inicien con "<<external_doc_"
                                  con el valor siendo una ruta (string) o lista de rutas a 
                                  documentos Word externos.
    
    Comportamiento:
        - Busca marcadores "<<external_doc_*>>" en los párrafos
        - Inserta el contenido completo de cada documento externo (párrafos y tablas)
        - Copia imágenes y mantiene relaciones correctas
        - Excluye headers, footers, marcas de agua y configuraciones de sección
        - Agrega saltos de página entre múltiples documentos
    
    Ejemplo de uso:
        from docx import Document
        from word_template_writer import insertar_documento_externo_en_plantilla
        
        doc = Document('plantilla.docx')
        diccionario = {
            "<<external_doc_plan_de_crucero>>": "plan_crucero.docx",
            # O múltiples documentos:
            # "<<external_doc_plan_de_crucero>>": ["plan1.docx", "plan2.docx"],
        }
        
        insertar_documento_externo_en_plantilla(doc, diccionario)
        doc.save('documento_con_plan.docx')
    """
     # Filtrar solo variables que NO son figuras, referencias o documentos externos
    variables_texto = {k: v for k, v in diccionario_de_reemplazos.items() if k.startswith("<<external_doc_")}
    
    # Para cada párrafo, procesar TODAS las variables de texto de una sola vez
    for parrafo in doc.paragraphs:
        # Encontrar todas las variables que están en este párrafo
        variables_en_parrafo = []
        for variable, dato in variables_texto.items():
            if variable in parrafo.text:
                variables_en_parrafo.append((variable, dato))        
                insert_external_document(parrafo, variable, dato, doc)
            
                msg = f"Documento insertado {variable}"
                if isinstance(dato, list) and len(dato) > 1:
                    msg = "Planes de crucero insertados"
                print(msg)


def rellenar_tablas_en_plantilla(doc, diccionario_de_reemplazos: dict):      
    """Rellena tablas en la plantilla usando keys <<editartabla...>>.
    
    Soporta estilos configurables, colores de fondo, MultiIndex, y merge de celdas.
    
    Args:
        doc: Objeto Document de python-docx
        diccionario_de_reemplazos: Diccionario con entradas que inician con <<editartabla
            Estructura por key:
            {
                "tabla": pd.DataFrame,
                "estilos_de_tabla": EstilosTabla | dict | None,
                "opciones_de_tabla": OpcionesTabla | dict | None,
            }
    
    Comportamiento:
        - Busca el marcador en cualquier celda de las tablas del documento
        - Inserta datos del DataFrame fila por fila, respetando orden de columnas
        - Aplica estilos según jerarquía: celda > fila > columna > defecto
        - Soporta colores de fondo RGB y estilos de párrafo del documento
        - Aplana MultiIndex automáticamente si está activado en opciones
        - Detecta y combina celdas verticales con valores repetidos (si está activado)
        - Elimina fila marcador automáticamente (si está activado)
    
    Ejemplo de uso:
        >>> import pandas as pd
        >>> from docx import Document
        >>> from word_template_writer import EstilosTabla, OpcionesTabla, rellenar_tablas_en_plantilla
        >>>
        >>> doc = Document('plantilla.docx')
        >>> estilos = EstilosTabla(doc)
        >>> opciones = OpcionesTabla()
        >>> df = pd.DataFrame({"Secuencia": [0, 1], "Localización": ["BOT-01", "BOT-02"]})
        >>>
        >>> diccionario = {
        ...     "<<editartabla_posiciones>>": {
        ...         "tabla": df,
        ...         "estilos_de_tabla": estilos,
        ...         "opciones_de_tabla": opciones,
        ...     }
        ... }
        >>>
        >>> rellenar_tablas_en_plantilla(doc, diccionario)
        >>> doc.save('documento_con_tablas_editadas.docx')
    
    Raises:
        ValueError: Si el marcador no se encuentra, si hay errores en la configuración,
                   o si los datos están vacíos
    
    Nota:
        - La plantilla debe tener un marcador <<nombre>> en alguna celda de la tabla
        - El marcador se coloca típicamente en la primera fila de datos (no encabezado)
        - Los datos se insertan por ORDEN de columna, no por nombre
        - Columna 0 del DataFrame → Columna 0 de la tabla
    """
    
    variables_texto = {
        k: v for k, v in diccionario_de_reemplazos.items()
        if k.startswith("<<editartabla")
    }
    
    for variable in variables_texto.keys():
        df_tabla = diccionario_de_reemplazos[variable]["tabla"]
        config_estilos = diccionario_de_reemplazos[variable]["estilos_de_tabla"]
        opciones_tabla = diccionario_de_reemplazos[variable]["opciones_de_tabla"]
        fill_table(doc, variable, df_tabla, config_estilos, opciones_tabla)
    
        msg = f"Tabla '{variable}' rellenada correctamente."
        print(msg)


def reemplazar_variables_en_tablas(doc, diccionario_de_reemplazos):
    """Reemplaza marcadores de posición en celdas de tablas del documento.
    
    Complementa a reemplazar_texto_en_plantilla() para tablas semi-estáticas
    donde solo se necesita reemplazar variables individuales, NO llenar
    toda la tabla con un DataFrame.
    
    Args:
        doc: Objeto Document de python-docx
        diccionario_de_reemplazos: Dict donde keys son marcadores (ej: "<<serial>>")
                                  y values son los textos de reemplazo
    
    Comportamiento:
        - Busca en TODAS las tablas del documento automáticamente
        - Ignora variables que inician con: <<fig, <<reffigura, <<reftabla_, <<nuevatabla_, <<editartabla, <<external_doc
        - Preserva el formato del texto original
        - Reutiliza la misma lógica de replace_text_variables_in_paragraph()
    
    Caso de uso típico:
        Tabla con:
        - Textos fijos en columna 0: "Serial", "Profundidad máxima", "Precisión"
        - Variables en columna 1: <<serial>>, <<profundidad>>, <<precision>>
        - Encabezado con variable: <<nombre_equipo>>
    
    Ejemplo:
        from docx import Document
        from word_template_writer import reemplazar_variables_en_tablas
        
        doc = Document('plantilla.docx')
        diccionario = {
            "<<nombre_equipo>>": "Sonda CTD #1",
            "<<serial>>": "4878505",
            "<<profundidad>>": "200 m",
            "<<precision>>": "±0.01°C"
        }
        
        reemplazar_variables_en_tablas(doc, diccionario)
        doc.save('resultado.docx')
    
    Ver también:
        - reemplazar_texto_en_plantilla(): para párrafos del documento (NO tablas)
        - rellenar_tablas_en_plantilla(): para tablas dinámicas con DataFrames
    """
    replace_text_variables_in_tables(doc, diccionario_de_reemplazos)
