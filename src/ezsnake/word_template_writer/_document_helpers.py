"""
Internal helper functions for inserting external Word documents.

This module contains low-level functions for merging external Word documents
into a template, including handling of images, watermarks, and relationships.

Private module - Not intended for direct external use.
Import from the public API in api.py instead.
"""

from copy import deepcopy
from genericpath import exists
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


def insert_external_document(paragraph, key, lista_rutas_documentos, doc_destino):
    """Inserta el contenido de uno o varios documentos Word externos.
    
    Args:
        paragraph: El párrafo donde está el marcador.
        key: El marcador a buscar.
        lista_rutas_documentos: Lista de rutas a documentos Word o una sola ruta (string).
        doc_destino: El documento de destino (Document) donde se insertará el contenido.
                     Necesario para copiar las relaciones de imágenes correctamente.
    
    Returns:
        True si se insertaron los documentos, False si no se encontró el marcador.
    
    Nota:
        Esta función solo copia párrafos y tablas, excluyendo:
        - Configuraciones de sección (w:sectPr) que incluyen headers/footers
        - Marcas de agua y fondos
        - Propiedades de documento
    """
    full_text = "".join(run.text for run in paragraph.runs)
    
    if key not in full_text:
        return False
    
    # Convertir a lista si es un solo string
    if isinstance(lista_rutas_documentos, str):
        lista_rutas_documentos = [lista_rutas_documentos]
    
    # Limpiar el párrafo marcador
    for run in paragraph.runs:
        run.text = ""
    
    # Obtener referencias
    p_element = paragraph._element
    parent = p_element.getparent()
    indice_insercion = parent.index(p_element)
    offset = 0
    
    # Insertar cada documento
    for ruta_doc in lista_rutas_documentos:
        if not exists(ruta_doc):
            # Agregar mensaje de error como párrafo
            p_error = OxmlElement('w:p')
            run = OxmlElement('w:r')
            text = OxmlElement('w:t')
            text.text = f"[Error: No se encontró {ruta_doc}]"
            run.append(text)
            p_error.append(run)
            parent.insert(indice_insercion + offset + 1, p_error)
            offset += 1
            continue
        
        # Cargar documento externo
        doc_externo = Document(ruta_doc)
        
        # Crear un mapeo de IDs de relación antiguo -> nuevo para las imágenes
        mapeo_rids = _copiar_relaciones_imagenes(doc_externo, doc_destino)
        
        # Copiar SOLO párrafos y tablas (excluir sectPr y otros elementos)
        for elemento in doc_externo.element.body:
            tag = elemento.tag
            
            # Solo copiar párrafos y tablas, ignorar el resto
            if tag == qn('w:p') or tag == qn('w:tbl'):
                # Verificar que no sea un párrafo con marca de agua
                if tag == qn('w:p') and _is_watermark(elemento):
                    continue
                
                elemento_copiado = deepcopy(elemento)
                
                # Actualizar los IDs de relación de imágenes en el elemento copiado
                _update_rids_in_element(elemento_copiado, mapeo_rids)
                
                parent.insert(indice_insercion + offset + 1, elemento_copiado)
                offset += 1
        
        # Agregar salto de página entre documentos (opcional)
        if lista_rutas_documentos.index(ruta_doc) < len(lista_rutas_documentos) - 1:
            # Agregar salto de página
            p_break = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            pageBreak = OxmlElement('w:pageBreakBefore')
            pPr.append(pageBreak)
            p_break.append(pPr)
            parent.insert(indice_insercion + offset + 1, p_break)
            offset += 1
    
    # Eliminar el párrafo marcador original
    parent.remove(p_element)
    
    return True


def _is_watermark(elemento_p):
    """Detecta si un párrafo es parte de una marca de agua.
    
    Args:
        elemento_p: Elemento XML de tipo w:p (párrafo)
    
    Returns:
        True si el párrafo contiene una marca de agua, False en caso contrario.
    """
    # Las marcas de agua suelen estar en párrafos con anchors de formas
    for run in elemento_p.findall(qn('w:r')):
        # Buscar elementos de dibujo (drawing)
        for drawing in run.findall(qn('w:drawing')):
            # Las marcas de agua tienen atributos específicos en sus anchors
            anchors = drawing.findall('.//' + qn('wp:anchor'))
            for anchor in anchors:
                # Verificar si tiene atributos típicos de marca de agua
                behindDoc = anchor.get(qn('behindDoc'))
                if behindDoc == '1':  # Marca de agua está detrás del texto
                    return True
    
    return False


def _copiar_relaciones_imagenes(doc_origen, doc_destino):
    """Copia las relaciones de imágenes del documento origen al destino.
    
    Args:
        doc_origen: Document de donde se copian las imágenes
        doc_destino: Document al que se copiarán las relaciones
    
    Returns:
        dict: Mapeo de rId antiguo -> rId nuevo para actualizar las referencias
              Ejemplo: {'rId1': 'rId15', 'rId2': 'rId16'}
    
    Esta función es necesaria para que las imágenes se vean correctamente
    en el documento destino después de copiar los elementos XML.
    """
    mapeo_rids = {}
    
    # Obtener las partes de relaciones
    part_origen = doc_origen.part
    part_destino = doc_destino.part
    
    # Iterar sobre todas las relaciones del documento origen
    for rel_id_antiguo, rel in part_origen.rels.items():
        # Solo copiar relaciones de imágenes
        if "image" in rel.reltype:
            try:
                # Obtener el blob de la imagen
                image_part = rel.target_part
                image_blob = image_part.blob
                content_type = image_part.content_type
                
                # Determinar la extensión del archivo según el content type
                ext = content_type.split('/')[-1]
                if ext == 'jpeg':
                    ext = 'jpg'
                
                # Crear una nueva imagen en el documento destino
                from io import BytesIO
                image_stream = BytesIO(image_blob)
                
                # Agregar la imagen al documento destino usando add_picture de manera indirecta
                # Crear un párrafo temporal para agregar la imagen
                temp_para = doc_destino.add_paragraph()
                temp_run = temp_para.add_run()
                picture = temp_run.add_picture(image_stream, width=Inches(1))
                
                # Obtener el rId de la imagen recién agregada
                inline = temp_run._element.find('.//' + qn('a:blip'))
                if inline is not None:
                    rel_id_nuevo = inline.get(qn('r:embed'))
                    mapeo_rids[rel_id_antiguo] = rel_id_nuevo
                
                # Eliminar el párrafo temporal
                p_element = temp_para._element
                p_element.getparent().remove(p_element)
                    
            except Exception as e:
                # Si hay error al copiar una imagen específica, continuar con las demás
                print(f"Advertencia: No se pudo copiar la imagen {rel_id_antiguo}: {str(e)}")
                continue
    
    return mapeo_rids


def _update_rids_in_element(elemento, mapeo_rids):
    """Actualiza los IDs de relación (rId) en un elemento XML con el mapeo proporcionado.
    
    Args:
        elemento: Elemento XML (párrafo o tabla copiado)
        mapeo_rids: Diccionario con el mapeo de rId antiguo -> rId nuevo
    
    Esta función busca todos los atributos r:embed y r:link en el elemento
    y los actualiza con los nuevos IDs.
    """
    if not mapeo_rids:
        return
    
    # Buscar todos los elementos con atributo r:embed (imágenes embebidas)
    for blip in elemento.findall('.//' + qn('a:blip')):
        rid_antiguo = blip.get(qn('r:embed'))
        if rid_antiguo and rid_antiguo in mapeo_rids:
            blip.set(qn('r:embed'), mapeo_rids[rid_antiguo])
    
    # Buscar todos los elementos con atributo r:link (imágenes vinculadas)
    for blip in elemento.findall('.//' + qn('a:blip')):
        rid_antiguo = blip.get(qn('r:link'))
        if rid_antiguo and rid_antiguo in mapeo_rids:
            blip.set(qn('r:link'), mapeo_rids[rid_antiguo])
