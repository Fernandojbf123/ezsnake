"""
Internal helper functions for text manipulation in Word templates.

This module contains low-level functions for replacing text variables in paragraphs.

Private module - Not intended for direct external use.
Import from the public API in api.py instead.
"""


def _parse_item_lista_a_texto_y_estilo(item):
    """Normaliza un elemento de lista a (texto, estilo).

    Formatos soportados:
        - "texto"
        - ("texto", "NombreEstilo")
        - ("texto", "") -> usa estilo Normal
    """
    if isinstance(item, tuple):
        texto = str(item[0]) if len(item) > 0 else ""
        estilo = str(item[1]).strip() if len(item) > 1 and item[1] is not None else ""
        return texto, estilo

    return str(item), ""


def replace_text_variables_in_paragraph(paragraph, lista_variables):
    """Reemplaza múltiples marcadores de posición en un párrafo de Word de una sola vez.
    
    Args:
        paragraph: El párrafo donde buscar los marcadores.
        lista_variables: Lista de tuplas (key, value) con los marcadores y sus valores.
                        Ejemplo: [("<<orden_de_servicio>>", "202"), ("<<numero_de_sondas>>", 5)]
    
    Esta función reemplaza todas las variables en un solo paso, evitando problemas
    de estado cuando hay múltiples variables en el mismo párrafo.
    
    Casos especiales:
        - Si el valor es una lista y el marcador es la única variable en el párrafo,
                    se crearán múltiples párrafos (uno por cada elemento de la lista).
                    Cada elemento puede ser:
                        * "texto" (usa estilo Normal)
                        * ("texto", "NombreEstilo")
                        * ("", "") para insertar salto de línea con estilo Normal
        - Si hay múltiples variables en el párrafo o hay texto adicional,
          solo se usa el primer elemento de la lista.
    """
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph
    
    # Obtener el texto completo del párrafo
    full_text = "".join(run.text for run in paragraph.runs)
    
    # Caso especial: Una sola variable con valor tipo lista que ocupa todo el párrafo
    if len(lista_variables) == 1:
        key, value = lista_variables[0]
        # Verificar si es una lista y el marcador ocupa todo el párrafo
        if isinstance(value, list) and len(value) > 0 and full_text.strip() == key:
            # Copiar el estilo del párrafo original
            estilo_original = paragraph.style
            estilo_normal = "Normal"

            # Parsear el primer elemento de la lista (texto + estilo opcional)
            texto_inicial, estilo_inicial = _parse_item_lista_a_texto_y_estilo(value[0])
            estilo_a_aplicar = estilo_inicial if estilo_inicial else estilo_normal

            # Aplicar estilo del primer elemento; fallback a estilo original si no existe
            try:
                paragraph.style = estilo_a_aplicar
            except KeyError:
                paragraph.style = estilo_original
            
            # Modificar el primer párrafo con el primer elemento
            primer_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            primer_run.text = texto_inicial
            for i in range(1, len(paragraph.runs)):
                paragraph.runs[i].text = ""
            
            # Obtener el elemento actual y su padre
            elemento_actual = paragraph._element
            
            # Insertar párrafos adicionales para el resto de elementos
            for item in value[1:]:
                texto, estilo = _parse_item_lista_a_texto_y_estilo(item)
                # Crear un nuevo elemento de párrafo
                nuevo_elemento = OxmlElement('w:p')
                
                # Insertar el nuevo párrafo después del anterior
                elemento_actual.addnext(nuevo_elemento)
                
                # Crear objeto Paragraph desde el elemento XML
                nuevo_parrafo = Paragraph(nuevo_elemento, paragraph._parent)
                estilo_a_aplicar = estilo if estilo else estilo_normal
                try:
                    nuevo_parrafo.style = estilo_a_aplicar
                except KeyError:
                    nuevo_parrafo.style = estilo_original
                
                # Agregar el texto al nuevo párrafo
                nuevo_parrafo.add_run(texto)
                
                # Actualizar el elemento actual para la siguiente iteración
                elemento_actual = nuevo_elemento
            
            return paragraph
    
    # Comportamiento estándar: reemplazar variables preservando formato de runs
    
    # Crear diccionario de reemplazos
    reemplazos = {}
    for key, value in lista_variables:
        if isinstance(value, list) and len(value) > 0:
            new_value, _ = _parse_item_lista_a_texto_y_estilo(value[0])
        else:
            new_value = str(value)
        reemplazos[key] = new_value
    
    # Estrategia 1: Intentar reemplazar run por run (caso óptimo - preserva formato)
    for run in paragraph.runs:
        texto_run = run.text
        for key, new_value in reemplazos.items():
            if key in texto_run:
                texto_run = texto_run.replace(key, new_value)
        run.text = texto_run
    
    # Verificar si aún quedan marcadores sin reemplazar (estaban partidos entre runs)
    texto_actual = "".join(run.text for run in paragraph.runs)
    marcadores_pendientes = [key for key in reemplazos.keys() if key in texto_actual]
    
    if not marcadores_pendientes:
        # Todos los reemplazos se hicieron en la Estrategia 1
        return paragraph
    
    # Estrategia 2: Reconstruir preservando formato (para marcadores partidos entre runs)
    # Guardar información de formato de cada run con su posición en el texto
    runs_info = []
    pos = 0
    for run in paragraph.runs:
        run_len = len(run.text)
        runs_info.append({
            'start': pos,
            'end': pos + run_len,
            'text': run.text,
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'font_name': run.font.name if run.font.name else None,
            'font_size': run.font.size,
            'color': run.font.color.rgb if run.font.color.rgb else None,
        })
        pos += run_len
    
    # Hacer los reemplazos en el texto completo
    texto_nuevo = texto_actual
    for key, new_value in reemplazos.items():
        texto_nuevo = texto_nuevo.replace(key, new_value)
    
    # Si no hubo cambios, retornar (no debería pasar, pero por seguridad)
    if texto_nuevo == texto_actual:
        return paragraph
    
    # Calcular el mapeo de posiciones: posición_nueva -> formato_original
    # Esto es complejo, así que usaremos una aproximación: mapear por carácter
    formato_por_posicion = []
    for run_info in runs_info:
        for i in range(run_info['start'], run_info['end']):
            formato_por_posicion.append({
                'bold': run_info['bold'],
                'italic': run_info['italic'],
                'underline': run_info['underline'],
                'font_name': run_info['font_name'],
                'font_size': run_info['font_size'],
                'color': run_info['color'],
            })
    
    # Limpiar todos los runs del párrafo
    for run in paragraph.runs:
        run.text = ""
    
    # Reconstruir el párrafo aplicando formato según el texto original
    # Usamos una heurística: aplicar el formato del primer carácter de cada segmento
    pos_original = 0
    pos_nueva = 0
    
    while pos_nueva < len(texto_nuevo):
        # Encontrar hasta dónde llega este segmento con el mismo formato
        if pos_original < len(formato_por_posicion):
            formato_actual = formato_por_posicion[pos_original]
        else:
            # Si nos pasamos del texto original (por reemplazos más largos), usar formato por defecto
            formato_actual = formato_por_posicion[-1] if formato_por_posicion else {}
        
        # Encontrar cuántos caracteres consecutivos tienen el mismo formato
        longitud_segmento = 1
        while (pos_nueva + longitud_segmento < len(texto_nuevo) and 
               pos_original + longitud_segmento < len(formato_por_posicion) and
               formato_por_posicion[pos_original + longitud_segmento] == formato_actual):
            longitud_segmento += 1
        
        # Ajustar si el segmento es más largo debido a un reemplazo
        texto_segmento = texto_nuevo[pos_nueva:pos_nueva + longitud_segmento]
        
        # Crear un nuevo run con este texto y formato
        new_run = paragraph.add_run(texto_segmento)
        new_run.bold = formato_actual.get('bold')
        new_run.italic = formato_actual.get('italic')
        new_run.underline = formato_actual.get('underline')
        if formato_actual.get('font_name'):
            new_run.font.name = formato_actual.get('font_name')
        if formato_actual.get('font_size'):
            new_run.font.size = formato_actual.get('font_size')
        if formato_actual.get('color'):
            new_run.font.color.rgb = formato_actual.get('color')
        
        pos_nueva += longitud_segmento
        pos_original += longitud_segmento
    
    return paragraph


def replace_text_variables_in_tables(doc, diccionario_de_reemplazos):
    """Reemplaza marcadores de posición en todas las celdas de todas las tablas del documento.
    
    Args:
        doc: Objeto Document de python-docx
        diccionario_de_reemplazos: Dict con {<<variable>>: valor}
    
    Comportamiento:
        - Itera todas las tablas del documento
        - Busca en todos los párrafos de cada celda
        - Ignora marcadores que inician con: <<fig, <<reffigura, <<reftabla_, <<nuevatabla_, <<editartabla, <<external_doc
        - Reutiliza replace_text_variables_in_paragraph() para preservar formato
    
    Caso de uso típico:
        Tablas semi-estáticas con textos fijos + variables individuales.
        
        Ejemplo en plantilla Word:
        ┌──────────────────────────────┬─────────────────────┐
        │        <<nombre_equipo>>                           │ ← Encabezado
        ├──────────────────────────────┼─────────────────────┤
        │ Serial                       │ <<serial>>          │ ← Fijos + variables
        │ Profundidad máxima           │ <<profundidad>>     │
        └──────────────────────────────┴─────────────────────┘
    
    Nota:
        Complementa a reemplazar_texto_en_plantilla() que solo busca en doc.paragraphs.
        Para tablas dinámicas (llenar con DataFrame completo), usar rellenar_tablas_en_plantilla().
    """
    # Filtrar variables de texto (ignorar fig, reffigura, tabla, external_doc)
    # Verificar que INICIE con estos prefijos, no solo que los contenga
    prefijos_excluidos = [
        "<<fig",
        "<<reffigura",
        "<<reftabla_",
        "<<nuevatabla_",
        "<<editartabla",
        "<<external_doc",
    ]
    variables_texto = {
        k: v for k, v in diccionario_de_reemplazos.items() 
        if not any(k.startswith(prefix) for prefix in prefijos_excluidos)
    }
    
    # Iterar todas las tablas del documento
    for table in doc.tables:
        # Iterar todas las celdas de la tabla
        for row in table.rows:
            for cell in row.cells:
                # Cada celda puede tener múltiples párrafos
                for parrafo in cell.paragraphs:
                    # Encontrar variables en este párrafo
                    variables_en_parrafo = []
                    for variable, dato in variables_texto.items():
                        if variable in parrafo.text:
                            variables_en_parrafo.append((variable, dato))
                    
                    # Si hay variables, reemplazarlas (reutiliza función existente)
                    if variables_en_parrafo:
                        replace_text_variables_in_paragraph(parrafo, variables_en_parrafo)
    
    msg = "Se reemplazaron las variables en las tablas del documento."
    print(msg)
